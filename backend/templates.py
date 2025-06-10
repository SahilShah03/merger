from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import xml.etree.ElementTree as ET
import os
from datetime import datetime
import logging
import aiohttp
import asyncio

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

async def get_cve_details(cve_id):
    """Fetch CVE details from NVD API"""
    async with aiohttp.ClientSession() as session:
        url = f"https://services.nvd.nist.gov/rest/json/cves/2.0?cveId={cve_id}"
        async with session.get(url) as response:
            if response.status == 200:
                data = await response.json()
                return data
    return None

def get_owasp_category(plugin_id):
    """Map Nessus plugin IDs to OWASP categories"""
    owasp_categories = {
        "web": ["10020", "10021", "10022"],  # Example plugin IDs
        "injection": ["10023", "10024"],
        "authentication": ["10025", "10026"],
        "crypto": ["10027", "10028"],
        "config": ["10029", "10030"]
    }
    
    for category, plugins in owasp_categories.items():
        if plugin_id in plugins:
            return category
    return "Other"

def generate_report(nessus_file_path):
    try:
        tree = ET.parse(nessus_file_path)
        root = tree.getroot()
    except Exception:
        doc = Document()
        doc.add_heading('Vulnerability Assessment Report', 0)
        doc.add_paragraph('Error: Invalid or corrupt Nessus file.')
        output_path = os.path.join('/tmp/', 'vulnerability_report.docx')
        doc.save(output_path)
        return output_path

    doc = Document()
    doc.add_heading('Vulnerability Assessment Report', 0)
    doc.add_paragraph('This is a simple vulnerability report.')
    output_path = os.path.join('/tmp/', 'vulnerability_report.docx')
    doc.save(output_path)
    return output_path 